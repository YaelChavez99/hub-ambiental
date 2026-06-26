"""
Módulo: exportador_word.py
Exportación a Word (.docx) profesional del Capítulo 5 y la tabla NOM-138.

Genera documentos con:
  - Portada con datos del proyecto
  - Capítulo 5 completo con estilos de encabezado
  - Tabla NOM-138 con colores de excedencias
  - Pie de página con número de proyecto
  - Formato tamaño carta, márgenes 1 pulgada, fuente Arial 11pt

Compatible con: python-docx>=1.1, streamlit>=1.35
"""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any

import streamlit as st

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Cm, Pt, RGBColor, Inches
    _DOCX_DISPONIBLE = True
except ImportError:
    _DOCX_DISPONIBLE = False


# ---------------------------------------------------------------------------
# Colores oficiales del documento (coinciden con tabla_nom138.py)
# ---------------------------------------------------------------------------
_AZUL_HEADER  = RGBColor(0x1F, 0x38, 0x64)   # #1F3864
_CYAN_EXCEDE  = RGBColor(0x00, 0xB0, 0xF0)   # #00B0F0
_AMARILLO_LMP = RGBColor(0xFF, 0xFF, 0x99)   # #FFFF99
_GRIS_ZONA    = RGBColor(0xD9, 0xE1, 0xF2)   # #D9E1F2
_BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)
_NEGRO        = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Helpers de formato Word
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, color: RGBColor) -> None:
    """Aplica color de fondo a una celda de tabla en Word."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    # RGBColor es una tuple — str() devuelve directamente el hex (ej. 'D9E1F2')
    hex_color = str(color)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_border(cell) -> None:
    """Aplica borde fino a una celda."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFBFBF")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _celda_texto(
    cell, texto: str,
    bold: bool = False,
    font_size: int = 9,
    color_texto: RGBColor = _NEGRO,
    alinear: str = "center",
) -> None:
    """Escribe texto en una celda con formato."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if alinear == "center"
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    run = para.add_run(str(texto))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = color_texto
    run.font.name = "Arial"
    _set_cell_border(cell)


def _configurar_documento(doc: "Document") -> None:
    """Aplica configuración global: tamaño carta, márgenes, fuente."""
    from docx.shared import Inches
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)


def _aplicar_estilo_heading(doc: "Document") -> None:
    """Configura los estilos de encabezado para que sean consistentes."""
    from docx.shared import Pt, RGBColor
    styles_config = {
        "Heading 1": (14, True,  _AZUL_HEADER),
        "Heading 2": (12, True,  _AZUL_HEADER),
        "Heading 3": (11, True,  _NEGRO),
    }
    for nombre, (size, bold, color) in styles_config.items():
        try:
            style = doc.styles[nombre]
            style.font.name  = "Arial"
            style.font.size  = Pt(size)
            style.font.bold  = bold
            style.font.color.rgb = color
        except KeyError:
            pass


def _agregar_pie_pagina(doc: "Document", texto: str) -> None:
    """Agrega pie de página con texto del proyecto y número de página."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    for section in doc.sections:
        footer = section.footer
        para   = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"{texto} | Página ")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        # Número de página automático
        fldChar_b = OxmlElement("w:fldChar")
        fldChar_b.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = " PAGE "
        fldChar_e = OxmlElement("w:fldChar")
        fldChar_e.set(qn("w:fldCharType"), "end")
        run_num = para.add_run()
        run_num.font.name = "Arial"
        run_num.font.size = Pt(8)
        run_num.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_num._r.append(fldChar_b)
        run_num._r.append(instrText)
        run_num._r.append(fldChar_e)


# ---------------------------------------------------------------------------
# Generador principal: Capítulo 5 en Word
# ---------------------------------------------------------------------------

def generar_word_capitulo5(
    texto_cap5:        str,
    id_proyecto:       str,
    nombre_siniestro:  str,
    uso_suelo:         str,
    municipio:         str = "",
    estado:            str = "",
    responsable:       str = "",
    datos_inegi:       dict | None = None,
    datos_conagua:     dict | None = None,
) -> bytes | None:
    """
    Genera el Capítulo 5 completo como archivo .docx profesional.

    Args:
        texto_cap5: Texto markdown del capítulo (salida de generar_capitulo5_claude)
        id_proyecto: ID del proyecto para el pie de página
        nombre_siniestro: Para la portada
        ...

    Returns:
        bytes del archivo .docx listo para descargar, o None si falla.
    """
    if not _DOCX_DISPONIBLE:
        st.error("python-docx no está instalado. Ejecuta: pip install python-docx")
        return None

    if not texto_cap5:
        st.error("No hay texto del Capítulo 5 para exportar.")
        return None

    doc = Document()
    _configurar_documento(doc)
    _aplicar_estilo_heading(doc)
    _agregar_pie_pagina(doc, f"Proyecto {id_proyecto}")

    # ── PORTADA ────────────────────────────────────────────────────────────
    doc.add_paragraph()   # espacio superior
    doc.add_paragraph()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titulo.add_run("INFORME DE CARACTERIZACIÓN DE SITIO CONTAMINADO")
    run_t.bold = True
    run_t.font.size = Pt(16)
    run_t.font.name = "Arial"
    run_t.font.color.rgb = _AZUL_HEADER

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = subtitulo.add_run(nombre_siniestro.upper() if nombre_siniestro else "")
    run_s.font.size = Pt(13)
    run_s.font.name = "Arial"
    run_s.font.color.rgb = _AZUL_HEADER

    doc.add_paragraph()

    # Tabla de datos del proyecto en portada
    tabla_port = doc.add_table(rows=6, cols=2)
    tabla_port.style = "Table Grid"
    datos_portada = [
        ("Número de proyecto:",  id_proyecto),
        ("Municipio / Estado:",  f"{municipio}, {estado}" if municipio else estado),
        ("Uso de suelo:",        uso_suelo),
        ("Responsable técnico:", responsable or "—"),
        ("Fecha de generación:", datetime.now().strftime("%d de %B de %Y")),
        ("Normatividad:",        "NOM-138-SEMARNAT/SSA1-2012"),
    ]
    for i, (etiq, val) in enumerate(datos_portada):
        c_etiq = tabla_port.rows[i].cells[0]
        c_val  = tabla_port.rows[i].cells[1]
        _set_cell_bg(c_etiq, _GRIS_ZONA)
        _celda_texto(c_etiq, etiq, bold=True,  font_size=10, alinear="left")
        _celda_texto(c_val,  val,  bold=False, font_size=10, alinear="left")

    # Fuentes de datos verificados
    if datos_inegi or datos_conagua:
        doc.add_paragraph()
        para_fuentes = doc.add_paragraph()
        run_f = para_fuentes.add_run("Datos institucionales verificados:")
        run_f.bold = True
        run_f.font.size = Pt(9)
        run_f.font.name = "Arial"
        run_f.font.color.rgb = RGBColor(0x15, 0x57, 0x24)

        if datos_inegi and datos_inegi.get("origen") not in ("no_disponible", None):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(
                f"INEGI — Censo de Población y Vivienda 2020 · "
                f"Clave {datos_inegi.get('clave_entidad','')}"
                f"{datos_inegi.get('clave_municipio','')} · "
                f"Población: {datos_inegi.get('poblacion_total','—')} hab."
            ).font.size = Pt(9)

        if datos_conagua:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(
                f"CONAGUA — {datos_conagua.get('nombre','—')} "
                f"(clave {datos_conagua.get('clave','—')}) · "
                f"{datos_conagua.get('fuente','DOF/CONAGUA')}"
            ).font.size = Pt(9)

    doc.add_page_break()

    # ── CUERPO DEL CAPÍTULO ────────────────────────────────────────────────
    # Parsear el markdown del Capítulo 5 y convertir a párrafos Word
    lineas = texto_cap5.splitlines()

    for linea in lineas:
        linea_strip = linea.strip()
        if not linea_strip:
            doc.add_paragraph()
            continue

        # Encabezado nivel 1 (## 5. TÍTULO)
        if linea_strip.startswith("## "):
            texto_h = linea_strip[3:].strip()
            heading = doc.add_heading(texto_h, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # Encabezado nivel 2 (### 5.1 Subtítulo)
        if linea_strip.startswith("### "):
            texto_h = linea_strip[4:].strip()
            heading = doc.add_heading(texto_h, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # Encabezado nivel 3 (#### Subtítulo menor)
        if linea_strip.startswith("#### "):
            texto_h = linea_strip[5:].strip()
            heading = doc.add_heading(texto_h, level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        # Lista con guión o asterisco
        if linea_strip.startswith(("- ", "* ")):
            texto_item = linea_strip[2:].strip()
            para = doc.add_paragraph(style="List Bullet")
            _agregar_runs_con_negrita(para, texto_item)
            continue

        # Párrafo normal — aplicar justificado y negrita inline
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        _agregar_runs_con_negrita(para, linea_strip)

    # ── SERIALIZAR A BYTES ─────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _agregar_runs_con_negrita(para, texto: str) -> None:
    """
    Parsea **texto en negrita** dentro de una línea y agrega runs
    con el formato correcto en el párrafo Word.
    """
    # Dividir por **...**
    partes = re.split(r'\*\*(.+?)\*\*', texto)
    for i, parte in enumerate(partes):
        if not parte:
            continue
        run = para.add_run(parte)
        run.font.name  = "Arial"
        run.font.size  = Pt(11)
        run.bold = (i % 2 == 1)   # los impares son los grupos capturados (negrita)


# ---------------------------------------------------------------------------
# Generador: Tabla NOM-138 en Word
# ---------------------------------------------------------------------------

def generar_word_tabla_nom138(
    historial:        list[dict],
    lim_vigentes:     dict[str, float],
    id_proyecto:      str,
    nombre_siniestro: str,
    uso_suelo:        str,
) -> bytes | None:
    """
    Genera la tabla NOM-138 completa como .docx con colores de excedencias.

    Args:
        historial:    Datos de cargar_laboratorio_proyecto()
        lim_vigentes: LMP por parámetro según uso de suelo
        id_proyecto:  ID del proyecto
        ...

    Returns:
        bytes del .docx listo para descargar, o None si falla.
    """
    if not _DOCX_DISPONIBLE:
        st.error("python-docx no está instalado. Ejecuta: pip install python-docx")
        return None

    if not historial:
        st.error("No hay muestras en el expediente para exportar.")
        return None

    def sf(v: Any) -> float:
        try:
            return float(str(v).replace(",", "")) if v else 0.0
        except (ValueError, TypeError):
            return 0.0

    doc = Document()
    _configurar_documento(doc)
    _agregar_pie_pagina(doc, f"NOM-138 · Proyecto {id_proyecto}")

    # ── Título ─────────────────────────────────────────────────────────────
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = titulo.add_run(nombre_siniestro.upper() if nombre_siniestro else "TABLA NOM-138")
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.name = "Arial"
    run_t.font.color.rgb = _AZUL_HEADER

    subtit = doc.add_paragraph()
    subtit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s  = subtit.add_run(
        f"NOM-138-SEMARNAT/SSA1-2012 · Uso de suelo: {uso_suelo} · "
        f"Proyecto: {id_proyecto}"
    )
    run_s.font.size = Pt(9)
    run_s.font.name = "Arial"
    run_s.font.color.rgb = _AZUL_HEADER

    doc.add_paragraph()

    # ── Parámetros a mostrar ───────────────────────────────────────────────
    PARAMS = ["HFL", "Benceno", "Tolueno", "Etilbenceno", "Xilenos"]
    ENCABEZADOS = [
        "ZONA AFECTADA", "MUESTRA", "PROF. (m)",
        "X (Este)", "Y (Norte)",
        "HFL", "BENCENO", "TOLUENO", "ETILBENCENO", "XILENOS",
        "pH", "HUM. (%)", "NOM-138",
    ]

    # Ordenar historial por zona y número de muestra
    import re as _re
    def sort_key(h):
        zona = str(h.get("zona", "")).upper()
        orden = {"ZONA 1": 0, "ZONA 2": 1, "ZONA 3": 2, "PERIFERIA": 3}
        muest = str(h.get("id_muestra", ""))
        nums  = _re.findall(r"[\d.]+", muest)
        n1 = float(nums[0]) if len(nums) > 0 else 0
        n2 = float(nums[1]) if len(nums) > 1 else 0
        return (orden.get(zona, 99), n1, n2)

    historial_sorted = sorted(historial, key=sort_key)

    # Tabla Word con orientación apaisada para caber todas las columnas
    tabla = doc.add_table(rows=1, cols=len(ENCABEZADOS))
    tabla.style = "Table Grid"

    # Anchos de columna en EMU (1 pulgada = 914400 EMU ≈ 9144 en DXA)
    # Ancho total disponible: 9 pulgadas (carta - márgenes)
    anchos_dxa = [800, 900, 600, 900, 900, 700, 700, 700, 800, 700, 500, 500, 700]

    # Fila de encabezado
    fila_enc = tabla.rows[0]
    for j, (enc, ancho) in enumerate(zip(ENCABEZADOS, anchos_dxa)):
        celda = fila_enc.cells[j]
        _set_cell_bg(celda, _AZUL_HEADER)
        _celda_texto(celda, enc, bold=True, font_size=7,
                     color_texto=_BLANCO)
        celda.width = Pt(ancho)

    # Filas de datos
    zona_actual = None
    for h in historial_sorted:
        res   = h.get("resultados", {})
        zona  = str(h.get("zona", "")).upper().strip()
        muest = str(h.get("id_muestra", ""))
        prof  = str(h.get("profundidad", ""))
        cx    = str(h.get("x", ""))
        cy    = str(h.get("y", ""))

        fila = tabla.add_row()

        # Celda zona — cambio de color si es nueva zona
        celda_zona = fila.cells[0]
        if zona != zona_actual:
            zona_actual = zona
            _set_cell_bg(celda_zona, _GRIS_ZONA)
            _celda_texto(celda_zona, zona, bold=True, font_size=8)
        else:
            _set_cell_bg(celda_zona, _GRIS_ZONA)
            _celda_texto(celda_zona, "", font_size=8)

        # Muestra, profundidad, coordenadas
        for j, val in enumerate([muest, prof, cx, cy], start=1):
            _celda_texto(fila.cells[j], val, font_size=8)

        # Parámetros analíticos con color por excedencia
        for j, param in enumerate(PARAMS, start=5):
            v_float = sf(res.get(param, 0))
            lmp     = lim_vigentes.get(param, 999999)
            if v_float <= 0:
                _celda_texto(fila.cells[j], "< L.C.", font_size=8)
            else:
                supera = v_float > lmp
                txt = (f"{v_float:,.2f}" if v_float >= 100
                       else f"{v_float:.2f}"  if v_float >= 10
                       else f"{v_float:.3f}")
                celda_p = fila.cells[j]
                if supera:
                    _set_cell_bg(celda_p, _CYAN_EXCEDE)
                _celda_texto(celda_p, txt, bold=supera, font_size=8)

        # pH y Humedad
        ph_v  = sf(res.get("pH",      0))
        hm_v  = sf(res.get("Humedad", 0))
        _celda_texto(fila.cells[10], f"{ph_v:.2f}"  if ph_v > 0 else "—", font_size=8)
        _celda_texto(fila.cells[11], f"{hm_v:.3f}"  if hm_v > 0 else "—", font_size=8)

        # Evaluación NOM-138
        celda_ev = fila.cells[12]
        if h.get("rebase"):
            _set_cell_bg(celda_ev, _CYAN_EXCEDE)
            _celda_texto(celda_ev, "EXCEDE", bold=True, font_size=8)
        else:
            _celda_texto(celda_ev, "CONFORME", font_size=8)

    # Fila de LMP
    fila_lmp = tabla.add_row()
    _set_cell_bg(fila_lmp.cells[0], _AMARILLO_LMP)
    _celda_texto(fila_lmp.cells[0], "LMP NOM-138", bold=True, font_size=7)
    _set_cell_bg(fila_lmp.cells[1], _AMARILLO_LMP)
    _celda_texto(fila_lmp.cells[1], uso_suelo, font_size=7)
    for j in range(2, 5):
        _set_cell_bg(fila_lmp.cells[j], _AMARILLO_LMP)
        _celda_texto(fila_lmp.cells[j], "—", font_size=7)
    for j, param in enumerate(PARAMS, start=5):
        celda_lmp = fila_lmp.cells[j]
        _set_cell_bg(celda_lmp, _AMARILLO_LMP)
        _celda_texto(celda_lmp, str(lim_vigentes.get(param, "—")),
                     bold=True, font_size=8)
    for j in range(10, 13):
        _set_cell_bg(fila_lmp.cells[j], _AMARILLO_LMP)
        _celda_texto(fila_lmp.cells[j], "—", font_size=7)

    # Nota al pie
    doc.add_paragraph()
    nota = doc.add_paragraph()
    run_n = nota.add_run(
        "< L.C.: Límite Cuantificable  ·  "
        "HFL=4.68  ·  Benceno=0.030  ·  Tolueno=0.10  ·  "
        "Etilbenceno=0.20  ·  m,p-Xilenos=0.30  ·  o-Xileno=0.20  mg/kg  ·  "
        "Valores resaltados en azul superan el LMP de la NOM-138-SEMARNAT/SSA1-2012"
    )
    run_n.font.size = Pt(8)
    run_n.font.name = "Arial"
    run_n.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Serializar
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Renderizado en Streamlit — botones de descarga Word
# ---------------------------------------------------------------------------

def render_descarga_word(
    texto_cap5:        str | None,
    historial_lab:     list[dict],
    lim_vigentes:      dict[str, float],
    id_proyecto:       str,
    nombre_siniestro:  str,
    uso_suelo:         str,
    municipio:         str = "",
    estado:            str = "",
    responsable:       str = "",
    datos_inegi:       dict | None = None,
    datos_conagua:     dict | None = None,
) -> None:
    """
    Renderiza los botones de descarga Word en la interfaz Streamlit.
    Llamar desde render_herramienta_dispersion() en la tab de Descargas.
    """
    if not _DOCX_DISPONIBLE:
        st.error(
            "❌ python-docx no está instalado. "
            "Agrega `python-docx>=1.1.0` en requirements.txt y reinicia."
        )
        return

    st.subheader("📄 Exportar a Word (.docx)")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**Capítulo 5 completo**")
        st.caption("Incluye portada, fuentes INEGI/CONAGUA verificadas y texto con estilos.")
        if texto_cap5:
            if st.button("⚙️ Generar Word — Capítulo 5", key="btn_word_cap5"):
                with st.spinner("Generando documento Word…"):
                    docx_bytes = generar_word_capitulo5(
                        texto_cap5, id_proyecto, nombre_siniestro, uso_suelo,
                        municipio, estado, responsable, datos_inegi, datos_conagua,
                    )
                if docx_bytes:
                    st.download_button(
                        "⬇️ Descargar Capítulo 5 (.docx)",
                        data=docx_bytes,
                        file_name=f"Cap5_{id_proyecto}.docx",
                        mime=(
                            "application/vnd.openxmlformats-officedocument"
                            ".wordprocessingml.document"
                        ),
                        use_container_width=True,
                    )
        else:
            st.info("Genera primero el Capítulo 5 en H4 — Dispersión.")

    with col2:
        st.markdown("**Tabla NOM-138 oficial**")
        st.caption("Tabla con colores de excedencias, fila LMP y nota de límites cuantificables.")
        if historial_lab:
            if st.button("⚙️ Generar Word — Tabla NOM-138", key="btn_word_nom138"):
                with st.spinner("Generando tabla NOM-138 en Word…"):
                    docx_bytes = generar_word_tabla_nom138(
                        historial_lab, lim_vigentes,
                        id_proyecto, nombre_siniestro, uso_suelo,
                    )
                if docx_bytes:
                    st.download_button(
                        "⬇️ Descargar Tabla NOM-138 (.docx)",
                        data=docx_bytes,
                        file_name=f"NOM138_{id_proyecto}.docx",
                        mime=(
                            "application/vnd.openxmlformats-officedocument"
                            ".wordprocessingml.document"
                        ),
                        use_container_width=True,
                    )
        else:
            st.info("Procesa primero el PDF de laboratorio en H3.")
